from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

URL_PATTERN = re.compile(r"(https?://[^\s<>()]+|www\.[^\s<>()]+)", re.IGNORECASE)
JLC_PCB_VENDOR_KEY = "jlc pcb"
JLC_PCB_EMAIL = "terpsracingev@gmail.com"
JLC_PCB_ENV_FILE = Path(__file__).resolve().parent.parent / "env" / "jlc_pcb.env"
JLC_PCB_PASSWORD_ENV = "JLC_PCB_PASSWORD"


def extract_url(value: Any) -> str:
    """Extract the first URL from the order form's long-text Link field."""
    if not isinstance(value, str):
        return ""
    text = value.strip()
    if not text:
        return ""
    match = URL_PATTERN.search(text)
    if match:
        url = match.group(1).rstrip(".,;:)]}")
        return f"https://{url}" if url.lower().startswith("www.") else url
    return text if not re.search(r"\s", text) else ""


def normalize_vendor(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip().casefold())


def jlc_pcb_password() -> str:
    load_dotenv(JLC_PCB_ENV_FILE, override=False)
    return os.getenv(JLC_PCB_PASSWORD_ENV, "")


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no public hyperlink helper, so build the Word XML node."""
    part = paragraph.part
    relationship_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)

    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_links_docx(vendor: str, items: list[dict]) -> bytes | None:
    """Build the vendor links document in memory for upload."""
    links = [(item.get("item", "Link"), extract_url(item.get("link", ""))) for item in items if extract_url(item.get("link", ""))]
    if not links:
        return None

    doc = Document()
    doc.add_paragraph(vendor, style="Title")
    if normalize_vendor(vendor) == JLC_PCB_VENDOR_KEY:
        doc.add_paragraph("Login credentials:")
        doc.add_paragraph(JLC_PCB_EMAIL)
        doc.add_paragraph(jlc_pcb_password() or f"Missing {JLC_PCB_PASSWORD_ENV}")

    for label, link in links:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_hyperlink(paragraph, link, str(label or link))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
